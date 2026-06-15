from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter


ROOT = Path("/Users/hang/Documents/Vx/backtests")
SOURCE = Path(
    "/Users/hang/Library/Containers/com.tencent.xinWeChat/Data/Documents/"
    "xwechat_files/wxid_49y1m38r00cm12_ca27/temp/RWTemp/2026-06/"
    "857c7b25ab63a620beaa8cc27be16e2f/Bloomberg Backtest.docx"
)


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
HEADER_FILL = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_doc(doc, subtitle=None):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Bloomberg Backtest Pack").font.color.rgb = MUTED

    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(subtitle)
        run.font.color.rgb = MUTED
        run.font.size = Pt(10)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)
    style_doc(doc, subtitle)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1.875, 4.625])
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for cell in hdr:
        set_cell_shading(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    return table


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        p.add_run(item)


def build_strategy_doc():
    doc = Document()
    add_title(
        doc,
        "Bloomberg Backtest Strategy Specification",
        "Temporary strategy document derived from the initial Bloomberg Backtest draft.",
    )
    doc.add_heading("Strategy Overview", level=1)
    add_kv_table(
        doc,
        [
            ("Strategy name", "Long-Only Equity Strategy"),
            ("Signal family", "Weekly multi-factor: technical trend + fundamental quality + earnings momentum"),
            ("Asset class", "Large-cap and mid-cap liquid equities; global or regional index components"),
            ("Decision bar", "Weekly, with daily confirmation"),
            ("Direction", "Long only"),
            ("Rebalance", "Weekly"),
            ("Execution", "Next weekly open after signal confirmation"),
        ],
    )

    doc.add_heading("Universe and Liquidity Filter", level=1)
    add_bullets(
        doc,
        [
            "Market capitalization greater than USD 5 billion.",
            "Average daily trading volume greater than 1,000,000 shares over the past 20 trading days.",
            "Exclude penny stocks; current price must be greater than USD 10.",
        ],
    )

    doc.add_heading("Required Entry Filters", level=1)
    add_matrix(
        doc,
        ["Filter group", "Rule", "Implementation note"],
        [
            ("Fundamental quality", "Trailing 12-month EPS > 0", "Must pass before ranking/entry."),
            ("Fundamental quality", "YoY EPS growth >= 15%", "Use latest reported period available in Bloomberg."),
            ("Fundamental quality", "YoY revenue growth >= 10%", "Use comparable reported period."),
            ("Fundamental quality", "ROE >= 12%", "Quality threshold."),
            ("Fundamental quality", "Debt-to-equity <= 0.8", "Leverage control."),
            ("Earnings momentum", "Latest actual EPS >= consensus estimate", "Positive or non-negative surprise."),
            ("Earnings momentum", "Consensus EPS estimate revised upward in the past 30 days", "Upward revisions must exceed downward revisions."),
            ("Analyst sentiment", "Net analyst rating >= Buy", "Bloomberg consensus rating bias."),
            ("Weekly technicals", "Weekly close > 40-week MA and > 13-week MA", "Long and medium trend confirmation."),
            ("Weekly technicals", "Weekly MACD line > signal line", "Bullish trend condition."),
            ("Weekly technicals", "Weekly RSI between 40 and 65", "Avoid weak and extremely overbought names."),
            ("Daily confirmation", "Daily RSI > 50 and no bearish engulfing pattern", "Short-term confirmation before weekly entry."),
        ],
        [1.45, 2.85, 2.20],
    )

    doc.add_heading("Entry, Sizing, and Risk Management", level=1)
    add_bullets(
        doc,
        [
            "Enter long at the next weekly open when all universe, quality, earnings, sentiment, and technical filters pass.",
            "Maximum single position size: 5% of NAV.",
            "Maximum total gross exposure: 60% of NAV.",
            "Maximum number of concurrent positions: 15.",
            "Initial stop loss: 2.0 x Weekly ATR(14) below entry price.",
            "No pyramiding or adding to existing positions.",
        ],
    )

    doc.add_heading("Exit Rules", level=1)
    doc.add_paragraph("Any trigger below causes a full exit at the next weekly open.")
    add_bullets(
        doc,
        [
            "Price hits the initial ATR stop loss.",
            "Weekly close falls below the 40-week moving average.",
            "Weekly MACD line crosses below the signal line.",
            "Latest quarterly EPS turns negative.",
            "Consensus EPS estimate is revised down significantly within 30 days.",
            "Weekly RSI exceeds 75, triggering extreme-overbought profit-taking.",
        ],
    )

    doc.add_heading("Bloomberg Backtest Settings", level=1)
    add_kv_table(
        doc,
        [
            ("Data period", "At least 10 years"),
            ("Transaction cost", "0.10% per trade"),
            ("Slippage", "0.05% per trade"),
            ("Dividends", "Reinvested"),
            ("Survivorship bias", "Adjusted"),
            ("Required outputs", "Total return, annualized return, Sharpe, max drawdown, win rate, profit factor, trades, average holding period, turnover"),
        ],
    )
    doc.save(ROOT / "backtest testing1 - strategy.docx")


def build_result_doc():
    doc = Document()
    add_title(
        doc,
        "Bloomberg Backtest Result Template",
        "Temporary result file for Bloomberg BT/BQL output. Result fields are intentionally blank until the backtest is run.",
    )
    doc.add_heading("Run Metadata", level=1)
    add_kv_table(
        doc,
        [
            ("Strategy", "Long-Only Equity Strategy"),
            ("Universe", "Global / regional large-cap and mid-cap liquid equities"),
            ("Backtest period", "Pending Bloomberg run"),
            ("Benchmark", "Pending selection"),
            ("Rebalance", "Weekly"),
            ("Costs", "0.10% transaction cost + 0.05% slippage per trade"),
            ("Data treatment", "Dividends reinvested; survivorship bias adjusted"),
        ],
    )
    doc.add_heading("Headline Results", level=1)
    add_matrix(
        doc,
        ["Metric", "Strategy", "Benchmark", "Notes"],
        [
            ("Total return", "Pending", "Pending", "Populate from Bloomberg BT output."),
            ("Annualized return", "Pending", "Pending", "Use CAGR over selected data period."),
            ("Sharpe ratio", "Pending", "Pending", "Confirm risk-free-rate assumption."),
            ("Maximum drawdown", "Pending", "Pending", "Use peak-to-trough drawdown."),
            ("Win rate", "Pending", "N/A", "Winning exited trades divided by total exited trades."),
            ("Profit factor", "Pending", "N/A", "Gross profit divided by gross loss."),
            ("Number of trades", "Pending", "N/A", "Completed entries/exits."),
            ("Average holding period", "Pending", "N/A", "Weeks."),
            ("Annual turnover", "Pending", "N/A", "Annualized portfolio turnover."),
        ],
        [1.55, 1.35, 1.35, 2.25],
    )
    doc.add_heading("Interpretation Checklist", level=1)
    add_bullets(
        doc,
        [
            "Confirm whether excess return survives transaction costs and slippage.",
            "Compare drawdown timing against benchmark drawdowns.",
            "Review whether concentration or sector exposure explains most performance.",
            "Test sensitivity to weekly RSI upper bound, ATR stop multiplier, and max gross exposure.",
            "Check whether the signal works consistently across regions rather than one market regime.",
        ],
    )
    doc.add_heading("Trade Review Fields", level=1)
    add_matrix(
        doc,
        ["Field", "Purpose", "Status"],
        [
            ("Top winners", "Identify repeated winning signal patterns.", "Pending"),
            ("Top losers", "Identify common failure modes and stop behavior.", "Pending"),
            ("Longest holds", "Check whether trend rules overstay mature winners.", "Pending"),
            ("Fast exits", "Check if daily confirmation is too loose.", "Pending"),
            ("Open positions at end", "Review mark-to-market exposure.", "Pending"),
        ],
        [1.55, 3.80, 1.15],
    )
    doc.add_heading("Temporary Conclusion", level=1)
    doc.add_paragraph(
        "No investment conclusion should be drawn from this result template until Bloomberg BT/BQL has populated the performance metrics, trade log, and benchmark comparison."
    )
    doc.save(ROOT / "backtest testing1 - result template.docx")


def build_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"
    header_fill = PatternFill("solid", fgColor=HEADER_FILL)
    thin = Side(style="thin", color="D9DEE7")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    rows = [
        ["Metric", "Strategy", "Benchmark", "Notes"],
        ["Total return", "", "", "Bloomberg BT output"],
        ["Annualized return", "", "", "CAGR"],
        ["Sharpe ratio", "", "", "Confirm risk-free-rate assumption"],
        ["Maximum drawdown", "", "", "Peak-to-trough"],
        ["Win rate", "", "N/A", "Exited trades"],
        ["Profit factor", "", "N/A", "Gross profit / gross loss"],
        ["Number of trades", "", "N/A", "Completed trades"],
        ["Average holding period (weeks)", "", "N/A", ""],
        ["Annual turnover", "", "N/A", ""],
    ]
    for row in rows:
        ws.append(row)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="0B2545")
        cell.fill = header_fill
    for row in ws.iter_rows():
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="center", wrap_text=True)
    widths = [30, 18, 18, 42]
    for i, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width

    trades = wb.create_sheet("Trade Log")
    trades.append(["Entry Date", "Exit Date", "Ticker", "Entry Price", "Exit Price", "Return", "Holding Weeks", "Exit Reason", "Notes"])
    for cell in trades[1]:
        cell.font = Font(bold=True, color="0B2545")
        cell.fill = header_fill
    for col, width in enumerate([14, 14, 14, 14, 14, 14, 16, 24, 40], 1):
        trades.column_dimensions[get_column_letter(col)].width = width
    for row in trades.iter_rows():
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="center", wrap_text=True)

    params = wb.create_sheet("Parameters")
    for row in [
        ["Parameter", "Value"],
        ["Market cap minimum", "USD 5 billion"],
        ["Average daily volume minimum", "1,000,000 shares over 20 trading days"],
        ["Price minimum", "USD 10"],
        ["EPS growth minimum", "15% YoY"],
        ["Revenue growth minimum", "10% YoY"],
        ["ROE minimum", "12%"],
        ["Debt-to-equity maximum", "0.8"],
        ["Weekly trend", "Close > 40W MA and > 13W MA"],
        ["Weekly RSI range", "40 to 65"],
        ["Max single position", "5% NAV"],
        ["Max gross exposure", "60% NAV"],
        ["Max positions", "15"],
        ["Initial stop", "2.0 x Weekly ATR(14) below entry"],
    ]:
        params.append(row)
    for cell in params[1]:
        cell.font = Font(bold=True, color="0B2545")
        cell.fill = header_fill
    params.column_dimensions["A"].width = 34
    params.column_dimensions["B"].width = 56
    for row in params.iter_rows():
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="center", wrap_text=True)

    wb.save(ROOT / "backtest testing1 - result template.xlsx")


def main():
    ROOT.mkdir(parents=True, exist_ok=True)
    copyfile(SOURCE, ROOT / "backtest testing1 - temporary source.docx")
    build_strategy_doc()
    build_result_doc()
    build_xlsx()


if __name__ == "__main__":
    main()
